"""Export tools for position description documents.

Provides export functionality to markdown and Word document formats.
Handles filename generation, path configuration, and document formatting.
"""

import os
import re
from pathlib import Path
from typing import Optional

from datetime import date

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from src.models.draft import DraftElement
from src.models.interview import InterviewData


# Default export directory
DEFAULT_EXPORT_PATH = Path("output")


def sanitize_filename(title: str) -> str:
    """
    Sanitize a position title for use as a filename.

    Removes or replaces characters that are unsafe for filenames
    and normalizes whitespace.

    Args:
        title: The position title or text to sanitize

    Returns:
        Sanitized string safe for use as filename
    """
    if not title:
        return "position_description"

    # Replace common problematic characters
    sanitized = title.lower()

    # Remove or replace unsafe characters
    # Keep alphanumeric, spaces, hyphens, underscores
    sanitized = re.sub(r"[^a-z0-9\s\-_]", "", sanitized)

    # Replace multiple spaces/whitespace with single underscore
    sanitized = re.sub(r"\s+", "_", sanitized)

    # Replace multiple underscores/hyphens with single
    sanitized = re.sub(r"[_\-]+", "_", sanitized)

    # Remove leading/trailing underscores
    sanitized = sanitized.strip("_")

    # Truncate to reasonable length (max 100 chars)
    if len(sanitized) > 100:
        sanitized = sanitized[:100].rsplit("_", 1)[0]

    # Fallback if empty
    if not sanitized:
        return "position_description"

    return sanitized


def generate_filename(
    interview_data: Optional[dict] = None,
    extension: str = ".md",
) -> str:
    """
    Generate a filename from interview data.

    Uses position title, series, and grade to create a descriptive filename.

    Args:
        interview_data: Optional interview data dict
        extension: File extension (default: ".md")

    Returns:
        Generated filename with extension
    """
    parts = []

    if interview_data:
        interview = InterviewData.model_validate(interview_data)

        # Add position title
        if interview.position_title.is_set and interview.position_title.value:
            parts.append(sanitize_filename(interview.position_title.value))

        # Add series-grade
        if interview.series.is_set and interview.series.value:
            series = interview.series.value
            grade = interview.grade.value if interview.grade.is_set else "TBD"
            parts.append(f"{series}_{grade}")

    if not parts:
        parts.append("position_description")

    filename = "_".join(parts)

    # Ensure extension starts with dot
    if not extension.startswith("."):
        extension = f".{extension}"

    return f"{filename}{extension}"


def get_export_path(
    filename: str,
    export_dir: Optional[Path] = None,
) -> Path:
    """
    Get the full export path for a file.

    Creates the export directory if it doesn't exist.

    Args:
        filename: The filename to export
        export_dir: Optional custom export directory (default: output/)

    Returns:
        Full path to the export file
    """
    if export_dir is None:
        export_dir = DEFAULT_EXPORT_PATH

    # Create directory if needed
    export_dir.mkdir(parents=True, exist_ok=True)

    return export_dir / filename


def export_to_markdown(
    draft_elements: list[dict],
    interview_data: Optional[dict] = None,
    export_dir: Optional[Path] = None,
    filename: Optional[str] = None,
) -> Path:
    """
    Export position description to markdown format.

    Assembles all draft elements into a formatted markdown document
    and writes to the specified path.

    Args:
        draft_elements: List of serialized DraftElement dicts
        interview_data: Optional interview data for header
        export_dir: Optional export directory (default: output/)
        filename: Optional custom filename (auto-generated if not provided)

    Returns:
        Path to the exported markdown file
    """
    from src.utils.document import assemble_final_document

    # Generate filename if not provided
    if filename is None:
        filename = generate_filename(interview_data, extension=".md")

    # Get full path
    export_path = get_export_path(filename, export_dir)

    # Assemble document content
    content = assemble_final_document(draft_elements, interview_data)

    # Write to file
    export_path.write_text(content, encoding="utf-8")

    return export_path


def export_to_word(
    draft_elements: list[dict],
    interview_data: Optional[dict] = None,
    export_dir: Optional[Path] = None,
    filename: Optional[str] = None,
) -> Path:
    """
    Export position description to Word document format (.docx).

    Creates a professionally formatted Word document with proper
    headings, paragraphs, and styling.

    Args:
        draft_elements: List of serialized DraftElement dicts
        interview_data: Optional interview data for header
        export_dir: Optional export directory (default: output/)
        filename: Optional custom filename (auto-generated if not provided)

    Returns:
        Path to the exported Word document
    """
    # Generate filename if not provided
    if filename is None:
        filename = generate_filename(interview_data, extension=".docx")

    # Get full path
    export_path = get_export_path(filename, export_dir)

    # Create Word document
    doc = Document()

    # Set up styles
    _setup_document_styles(doc)

    # Add header section
    _add_document_header(doc, interview_data)

    # Add each draft element as a section
    for element_dict in draft_elements:
        element = DraftElement.model_validate(element_dict)

        # Skip elements without content
        if not element.content:
            continue

        # Add section heading
        doc.add_heading(element.display_name, level=2)

        # Add content - handle markdown formatting
        _add_formatted_content(doc, element.content)

    # Save document
    doc.save(str(export_path))

    return export_path


def export_to_markdown_bytes(
    draft_elements: list[dict],
    interview_data: Optional[dict] = None,
) -> bytes:
    """Export position description to markdown as bytes (for API responses).

    Args:
        draft_elements: List of serialized DraftElement dicts
        interview_data: Optional interview data for header

    Returns:
        UTF-8 encoded markdown bytes
    """
    from src.utils.document import assemble_final_document

    content = assemble_final_document(draft_elements, interview_data)
    return content.encode("utf-8")


def export_to_word_bytes(
    draft_elements: list[dict],
    interview_data: Optional[dict] = None,
) -> bytes:
    """Export position description to Word document as bytes (for API responses).

    Args:
        draft_elements: List of serialized DraftElement dicts
        interview_data: Optional interview data for header

    Returns:
        Bytes of the .docx file
    """
    import io

    doc = Document()
    _setup_document_styles(doc)
    _add_document_header(doc, interview_data)

    for element_dict in draft_elements:
        element = DraftElement.model_validate(element_dict)
        if not element.content:
            continue
        doc.add_heading(element.display_name, level=2)
        _add_formatted_content(doc, element.content)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _setup_document_styles(doc: Document) -> None:
    """
    Configure document styles for professional formatting.

    Args:
        doc: The Word document to style
    """
    # Configure default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Configure heading styles
    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Calibri"
        heading_style.font.bold = True


def _add_document_header(doc: Document, interview_data: Optional[dict]) -> None:
    """
    Add the OF-8 style document header as a table with position metadata.

    Args:
        doc: The Word document
        interview_data: Optional interview data dict
    """
    # Title bar
    title = doc.add_heading("Position Description — OF-8", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if not interview_data:
        doc.add_paragraph()
        return

    interview = InterviewData.model_validate(interview_data)

    # Build a 5-column table matching OF-8 layout
    table = doc.add_table(rows=0, cols=5)
    table.style = "Table Grid"

    def _add_cell_content(cell, label: str, value: str | None) -> None:
        """Add a label + value pair to a table cell."""
        p = cell.paragraphs[0]
        run_label = p.add_run(f"{label}\n")
        run_label.bold = True
        run_label.font.size = Pt(8)
        run_val = p.add_run(value or "—")
        run_val.font.size = Pt(10)

    # Row 1: Title + Pay Plan / Series / Grade
    row1 = table.add_row()
    _merge_cells(row1, 0, 1)
    _add_cell_content(row1.cells[0], "1. Position Title",
                      interview.position_title.value if interview.position_title.is_set else None)
    _add_cell_content(row1.cells[2], "2. Pay Plan", "GS")
    _add_cell_content(row1.cells[3], "3. Series",
                      interview.series.value if interview.series.is_set else None)
    _add_cell_content(row1.cells[4], "4. Grade",
                      interview.grade.value if interview.grade.is_set else None)

    # Row 2: Organization
    row2 = table.add_row()
    _merge_cells(row2, 0, 4)
    org_val = None
    org_data = (interview.organization_hierarchy if interview.organization_hierarchy.is_set
                else interview.organization if interview.organization.is_set else None)
    if org_data and org_data.is_set:
        org = org_data.value
        org_val = " / ".join(org) if isinstance(org, list) else str(org)
    _add_cell_content(row2.cells[0], "5. Employing Department / Agency", org_val)

    # Row 3: Reports To + Supervisory + FLSA
    row3 = table.add_row()
    _merge_cells(row3, 0, 1)
    _add_cell_content(row3.cells[0], "6. Reports To (Title)",
                      interview.reports_to.value if interview.reports_to.is_set else None)
    sup_label = None
    if interview.is_supervisor.is_set:
        sup_label = "Supervisory" if interview.is_supervisor.value else "Non-Supervisory"
    _add_cell_content(row3.cells[2], "7. Supervisory Status", sup_label)
    _merge_cells(row3, 3, 4)
    flsa = None
    if interview.is_supervisor.is_set:
        flsa = "Exempt" if interview.is_supervisor.value else "Nonexempt"
    _add_cell_content(row3.cells[3], "8. FLSA Status", flsa)

    # Row 4: Admin fields
    row4 = table.add_row()
    grade_num = 0
    if interview.grade.is_set and interview.grade.value:
        try:
            grade_num = int(str(interview.grade.value).replace("GS-", "").strip())
        except ValueError:
            pass
    sensitivity = "Noncritical-Sensitive" if grade_num >= 13 else "Non-Sensitive" if grade_num > 0 else None
    _add_cell_content(row4.cells[0], "9. Position Sensitivity", sensitivity)
    _add_cell_content(row4.cells[1], "10. Competitive Level", None)
    _add_cell_content(row4.cells[2], "11. Position Number", None)
    _add_cell_content(row4.cells[3], "12. Classified By", None)
    _add_cell_content(row4.cells[4], "13. Date", date.today().strftime("%m/%d/%Y"))

    # Row 5: Supervisory details (conditional)
    if interview.is_supervisor.is_set and interview.is_supervisor.value:
        row5 = table.add_row()
        _merge_cells(row5, 0, 2)
        sup_detail = None
        if interview.supervised_employees.is_set:
            sup_val = interview.supervised_employees.value
            if isinstance(sup_val, list):
                sup_detail = "; ".join(str(s) for s in sup_val)
            elif isinstance(sup_val, dict):
                sup_detail = "; ".join(f"{k}: {v}" for k, v in sup_val.items())
            else:
                sup_detail = str(sup_val)
        _add_cell_content(row5.cells[0], "Employees Supervised", sup_detail)
        _merge_cells(row5, 3, 4)
        pct = None
        if interview.percent_supervising.is_set:
            pct = f"{interview.percent_supervising.value}%"
        _add_cell_content(row5.cells[3], "% Time Supervising", pct)

    doc.add_paragraph()


def _merge_cells(row, start: int, end: int) -> None:
    """Merge cells in a table row from start to end index (inclusive)."""
    row.cells[start].merge(row.cells[end])


def _add_formatted_content(doc: Document, content: str) -> None:
    """
    Add content to document with basic markdown formatting.

    Handles:
    - Bullet lists (- or *)
    - Numbered lists
    - Bold text (**text**)
    - Multiple paragraphs

    Args:
        doc: The Word document
        content: Markdown-formatted content string
    """
    lines = content.split("\n")
    current_paragraph_lines: list[str] = []

    for line in lines:
        stripped = line.strip()

        # Check for bullet list item
        if stripped.startswith(("- ", "* ", "• ")):
            # Flush current paragraph
            if current_paragraph_lines:
                _add_paragraph_with_formatting(doc, " ".join(current_paragraph_lines))
                current_paragraph_lines = []

            # Add list item
            bullet_text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_text_with_bold(p, bullet_text)

        # Check for numbered list item
        elif re.match(r"^\d+\.\s", stripped):
            # Flush current paragraph
            if current_paragraph_lines:
                _add_paragraph_with_formatting(doc, " ".join(current_paragraph_lines))
                current_paragraph_lines = []

            # Add numbered item
            num_text = re.sub(r"^\d+\.\s*", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_text_with_bold(p, num_text)

        # Empty line - new paragraph
        elif not stripped:
            if current_paragraph_lines:
                _add_paragraph_with_formatting(doc, " ".join(current_paragraph_lines))
                current_paragraph_lines = []

        # Regular text - accumulate
        else:
            current_paragraph_lines.append(stripped)

    # Flush remaining content
    if current_paragraph_lines:
        _add_paragraph_with_formatting(doc, " ".join(current_paragraph_lines))


def _add_paragraph_with_formatting(doc: Document, text: str) -> None:
    """
    Add a paragraph with bold formatting support.

    Args:
        doc: The Word document
        text: Text potentially containing **bold** markers
    """
    p = doc.add_paragraph()
    _add_text_with_bold(p, text)


def _add_text_with_bold(paragraph, text: str) -> None:
    """
    Add text to a paragraph, handling **bold** markers.

    Args:
        paragraph: The Word paragraph
        text: Text potentially containing **bold** markers
    """
    # Pattern to match **bold text**
    bold_pattern = re.compile(r"\*\*([^*]+)\*\*")

    last_end = 0
    for match in bold_pattern.finditer(text):
        # Add text before bold
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        # Add bold text
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])
