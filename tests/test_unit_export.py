"""Unit tests for export tools."""

import tempfile
from pathlib import Path

import pytest

from src.tools.export_tools import (
    DEFAULT_EXPORT_PATH,
    export_to_markdown,
    export_to_word,
    generate_filename,
    get_export_path,
    sanitize_filename,
)


# ============================================================================
# Test Data Fixtures
# ============================================================================


@pytest.fixture
def sample_interview_data():
    """Sample interview data for testing."""
    return {
        "position_title": {"value": "IT Specialist", "is_set": True},
        "series": {"value": "2210", "is_set": True},
        "grade": {"value": "13", "is_set": True},
        "organization": {
            "value": ["Department of Testing", "Office of Examples"],
            "is_set": True,
        },
        "reports_to": {"value": "Branch Chief", "is_set": True},
        "is_supervisor": {"value": False, "is_set": True},
        "major_duties": {"value": [], "is_set": False},
        "purpose": {"value": None, "is_set": False},
        "direct_reports": {"value": None, "is_set": False},
    }


@pytest.fixture
def sample_draft_elements():
    """Sample draft elements for testing."""
    return [
        {
            "name": "introduction",
            "display_name": "Introduction",
            "content": "This position serves as an **IT Specialist** responsible for systems administration.",
            "status": "approved",
            "revision_count": 0,
            "qa_passed": True,
        },
        {
            "name": "major_duties",
            "display_name": "Major Duties",
            "content": "- Administers IT systems\n- Provides technical support\n- Implements security measures",
            "status": "approved",
            "revision_count": 0,
            "qa_passed": True,
        },
        {
            "name": "factor_1_knowledge",
            "display_name": "Factor 1: Knowledge Required",
            "content": "1. Knowledge of operating systems\n2. Skill in network administration\n3. Ability to troubleshoot complex issues",
            "status": "approved",
            "revision_count": 0,
            "qa_passed": True,
        },
    ]


# ============================================================================
# sanitize_filename Tests
# ============================================================================


class TestSanitizeFilename:
    """Tests for sanitize_filename function."""

    def test_basic_sanitization(self):
        """Test basic title sanitization."""
        assert sanitize_filename("IT Specialist") == "it_specialist"

    def test_special_characters_removed(self):
        """Test removal of special characters."""
        assert sanitize_filename("IT/System Analyst!") == "itsystem_analyst"

    def test_multiple_spaces_collapsed(self):
        """Test multiple spaces become single underscore."""
        assert sanitize_filename("IT    Specialist") == "it_specialist"

    def test_mixed_characters(self):
        """Test complex string with mixed characters."""
        result = sanitize_filename("IT Specialist (Security) - Level III")
        assert result == "it_specialist_security_level_iii"

    def test_empty_string(self):
        """Test empty string returns default."""
        assert sanitize_filename("") == "position_description"

    def test_none_value(self):
        """Test None returns default."""
        assert sanitize_filename(None) == "position_description"

    def test_only_special_chars(self):
        """Test string with only special chars returns default."""
        assert sanitize_filename("@#$%^&*()") == "position_description"

    def test_long_title_truncated(self):
        """Test very long titles are truncated."""
        long_title = "This is a very long position title " * 10
        result = sanitize_filename(long_title)
        assert len(result) <= 100

    def test_preserves_numbers(self):
        """Test numbers are preserved."""
        assert sanitize_filename("GS-2210-13") == "gs_2210_13"

    def test_leading_trailing_underscores_stripped(self):
        """Test leading/trailing underscores removed."""
        assert sanitize_filename("  IT Specialist  ") == "it_specialist"


# ============================================================================
# generate_filename Tests
# ============================================================================


class TestGenerateFilename:
    """Tests for generate_filename function."""

    def test_with_full_interview_data(self, sample_interview_data):
        """Test filename generation with full data."""
        filename = generate_filename(sample_interview_data, extension=".md")
        assert filename == "it_specialist_2210_13.md"

    def test_with_docx_extension(self, sample_interview_data):
        """Test filename with .docx extension."""
        filename = generate_filename(sample_interview_data, extension=".docx")
        assert filename == "it_specialist_2210_13.docx"

    def test_extension_without_dot(self, sample_interview_data):
        """Test extension without leading dot."""
        filename = generate_filename(sample_interview_data, extension="md")
        assert filename == "it_specialist_2210_13.md"

    def test_without_interview_data(self):
        """Test filename without interview data."""
        filename = generate_filename(None, extension=".md")
        assert filename == "position_description.md"

    def test_with_empty_interview_data(self):
        """Test filename with empty interview data."""
        empty_data = {
            "position_title": {"value": None, "is_set": False},
            "series": {"value": None, "is_set": False},
            "grade": {"value": None, "is_set": False},
            "organization": {"value": None, "is_set": False},
            "reports_to": {"value": None, "is_set": False},
            "is_supervisor": {"value": None, "is_set": False},
            "major_duties": {"value": None, "is_set": False},
            "purpose": {"value": None, "is_set": False},
            "direct_reports": {"value": None, "is_set": False},
        }
        filename = generate_filename(empty_data, extension=".md")
        assert filename == "position_description.md"

    def test_with_grade_tbd(self):
        """Test filename when grade not set."""
        data = {
            "position_title": {"value": "Analyst", "is_set": True},
            "series": {"value": "0343", "is_set": True},
            "grade": {"value": None, "is_set": False},
            "organization": {"value": None, "is_set": False},
            "reports_to": {"value": None, "is_set": False},
            "is_supervisor": {"value": None, "is_set": False},
            "major_duties": {"value": None, "is_set": False},
            "purpose": {"value": None, "is_set": False},
            "direct_reports": {"value": None, "is_set": False},
        }
        filename = generate_filename(data, extension=".md")
        assert filename == "analyst_0343_TBD.md"


# ============================================================================
# get_export_path Tests
# ============================================================================


class TestGetExportPath:
    """Tests for get_export_path function."""

    def test_default_path(self):
        """Test default export path."""
        path = get_export_path("test.md")
        assert path == DEFAULT_EXPORT_PATH / "test.md"

    def test_custom_directory(self):
        """Test custom export directory."""
        with tempfile.TemporaryDirectory() as tmpdir:
            custom_dir = Path(tmpdir) / "exports"
            path = get_export_path("test.md", export_dir=custom_dir)
            assert path == custom_dir / "test.md"
            assert custom_dir.exists()

    def test_creates_directory(self):
        """Test that directory is created if it doesn't exist."""
        with tempfile.TemporaryDirectory() as tmpdir:
            new_dir = Path(tmpdir) / "new" / "nested" / "dir"
            assert not new_dir.exists()
            path = get_export_path("test.md", export_dir=new_dir)
            assert new_dir.exists()
            assert path == new_dir / "test.md"


# ============================================================================
# export_to_markdown Tests
# ============================================================================


class TestExportToMarkdown:
    """Tests for export_to_markdown function."""

    def test_basic_export(self, sample_draft_elements, sample_interview_data):
        """Test basic markdown export."""
        with tempfile.TemporaryDirectory() as tmpdir:
            export_dir = Path(tmpdir)
            path = export_to_markdown(
                sample_draft_elements,
                sample_interview_data,
                export_dir=export_dir,
            )

            assert path.exists()
            assert path.suffix == ".md"

            content = path.read_text()
            assert "# Position Description" in content
            assert "IT Specialist" in content
            assert "## Introduction" in content
            assert "## Major Duties" in content

    def test_custom_filename(self, sample_draft_elements, sample_interview_data):
        """Test export with custom filename."""
        with tempfile.TemporaryDirectory() as tmpdir:
            export_dir = Path(tmpdir)
            path = export_to_markdown(
                sample_draft_elements,
                sample_interview_data,
                export_dir=export_dir,
                filename="custom_name.md",
            )

            assert path.name == "custom_name.md"
            assert path.exists()

    def test_without_interview_data(self, sample_draft_elements):
        """Test export without interview data."""
        with tempfile.TemporaryDirectory() as tmpdir:
            export_dir = Path(tmpdir)
            path = export_to_markdown(
                sample_draft_elements,
                interview_data=None,
                export_dir=export_dir,
            )

            assert path.exists()
            content = path.read_text()
            assert "## Introduction" in content

    def test_empty_elements(self, sample_interview_data):
        """Test export with empty elements list."""
        with tempfile.TemporaryDirectory() as tmpdir:
            export_dir = Path(tmpdir)
            path = export_to_markdown(
                [],
                sample_interview_data,
                export_dir=export_dir,
            )

            assert path.exists()
            content = path.read_text()
            # Should contain header but no sections
            assert "No draft elements" in content


# ============================================================================
# export_to_word Tests
# ============================================================================


class TestExportToWord:
    """Tests for export_to_word function."""

    def test_basic_export(self, sample_draft_elements, sample_interview_data):
        """Test basic Word export."""
        with tempfile.TemporaryDirectory() as tmpdir:
            export_dir = Path(tmpdir)
            path = export_to_word(
                sample_draft_elements,
                sample_interview_data,
                export_dir=export_dir,
            )

            assert path.exists()
            assert path.suffix == ".docx"
            # Verify file is valid by checking size
            assert path.stat().st_size > 0

    def test_custom_filename(self, sample_draft_elements, sample_interview_data):
        """Test export with custom filename."""
        with tempfile.TemporaryDirectory() as tmpdir:
            export_dir = Path(tmpdir)
            path = export_to_word(
                sample_draft_elements,
                sample_interview_data,
                export_dir=export_dir,
                filename="custom_document.docx",
            )

            assert path.name == "custom_document.docx"
            assert path.exists()

    def test_without_interview_data(self, sample_draft_elements):
        """Test export without interview data."""
        with tempfile.TemporaryDirectory() as tmpdir:
            export_dir = Path(tmpdir)
            path = export_to_word(
                sample_draft_elements,
                interview_data=None,
                export_dir=export_dir,
            )

            assert path.exists()
            assert path.suffix == ".docx"

    def test_document_readable(self, sample_draft_elements, sample_interview_data):
        """Test that exported document can be read back."""
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            export_dir = Path(tmpdir)
            path = export_to_word(
                sample_draft_elements,
                sample_interview_data,
                export_dir=export_dir,
            )

            # Read the document back
            doc = Document(str(path))

            # Verify it has content
            assert len(doc.paragraphs) > 0

            # Check for expected text
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Position Description" in full_text
            assert "IT Specialist" in full_text

    def test_bullet_lists_converted(self, sample_interview_data):
        """Test that bullet lists are properly converted."""
        from docx import Document

        elements = [
            {
                "name": "test",
                "display_name": "Test Section",
                "content": "- Item one\n- Item two\n- Item three",
                "status": "approved",
                "revision_count": 0,
                "qa_passed": True,
            }
        ]

        with tempfile.TemporaryDirectory() as tmpdir:
            export_dir = Path(tmpdir)
            path = export_to_word(
                elements,
                sample_interview_data,
                export_dir=export_dir,
            )

            doc = Document(str(path))
            # Check that items are in the document
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Item one" in full_text
            assert "Item two" in full_text
            assert "Item three" in full_text

    def test_bold_text_formatting(self, sample_interview_data):
        """Test that bold markdown is converted."""
        from docx import Document

        elements = [
            {
                "name": "test",
                "display_name": "Test Section",
                "content": "This is **bold text** in a sentence.",
                "status": "approved",
                "revision_count": 0,
                "qa_passed": True,
            }
        ]

        with tempfile.TemporaryDirectory() as tmpdir:
            export_dir = Path(tmpdir)
            path = export_to_word(
                elements,
                sample_interview_data,
                export_dir=export_dir,
            )

            doc = Document(str(path))

            # Find paragraph with bold text
            found_bold = False
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.text == "bold text" and run.bold:
                        found_bold = True
                        break

            assert found_bold, "Bold text formatting not found"

    def test_empty_elements(self, sample_interview_data):
        """Test export with empty elements list."""
        with tempfile.TemporaryDirectory() as tmpdir:
            export_dir = Path(tmpdir)
            path = export_to_word(
                [],
                sample_interview_data,
                export_dir=export_dir,
            )

            # Should still create a valid document with header only
            assert path.exists()
            assert path.stat().st_size > 0
